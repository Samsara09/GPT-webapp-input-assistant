import os
import io
import tkinter as tk
from tkinter import filedialog, ttk
import PyPDF2
from PyPDF2 import PdfReader
from docx import Document
import requests
import mimetypes
from bs4 import BeautifulSoup

def read_pdf(file_path):
    pdf_file = PyPDF2.PdfReader(file_path)
    content = ""
    for page_num in range(len(pdf_file.pages)):
        content += pdf_file.pages[page_num].extract_text()
    return content

def read_html(content):
    soup = BeautifulSoup(content, 'lxml')
    text = soup.get_text(separator="\n")
    return text

def read_pdf_from_memory(content):
    with io.BytesIO(content) as pdf_stream:
        pdf_reader = PdfReader(pdf_stream)
        num_pages = len(pdf_reader.pages)
        text = ""

        for page_num in range(num_pages):
            page = pdf_reader.pages[page_num]
            text += page.extract_text()

    return text

def read_docx_from_memory(content):
    with io.BytesIO(content) as docx_stream:
        doc = Document(docx_stream)
        full_text = []

        for para in doc.paragraphs:
            full_text.append(para.text)

        return "\n".join(full_text)

def read_docx(file_path):
    doc = Document(file_path)
    content = ""
    for paragraph in doc.paragraphs:
        content += paragraph.text + "\n"
    return content

def read_txt(file_path):
    with open(file_path, 'r') as f:
        content = f.read()
    return content

def read_file(content, content_type):
    if content_type == "application/pdf":
        return read_pdf_from_memory(content)
    elif content_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
        return read_docx_from_memory(content)
    elif content_type == "text/plain":
        return content.decode("utf-8")
    elif content_type == "text/html":
        return read_html(content)
    else:
        print(f"Unsupported content type: {content_type}")
        return ""
    
def chunk_text(text, size):
    return [text[i:i+size] for i in range(0, len(text), size)]

def clear_chunks():
    chunk_listbox.delete(0, tk.END)
    chunk_texts.clear()

def open_file():
    clear_chunks()
    file_path = filedialog.askopenfilename(initialdir=os.getcwd(), title="Select a file",
                                            filetypes=(("All files", "*.*"), ("pdf files", "*.pdf"), ("docx files", "*.docx"), ("txt files", "*.txt")))
    if not file_path:
        return
    with open(file_path, "rb") as file:
        content = file.read()
    content_type = mimetypes.guess_type(file_path)[0]

    content = read_file(content, content_type)
    chunk_size = int(chunk_size_entry.get())
    chunks = chunk_text(content, chunk_size)

    for i, chunk in enumerate(chunks, start=1):
        chunk_listbox.insert(tk.END, f'Chunk {i}')
        chunk_texts.append(chunk)
       
def open_url():
    clear_chunks()
    url = url_entry.get().strip()
    if url:
        content, content_type = fetch_document_content(url)
        if content is None:
            return
    else:
        return

    content = read_file(content, content_type)
    chunk_size = int(chunk_size_entry.get())
    chunks = chunk_text(content, chunk_size)

    for i, chunk in enumerate(chunks, start=1):
        chunk_listbox.insert(tk.END, f'Chunk {i}')
        chunk_texts.append(chunk)

def fetch_document_content(url):
    try:
        response = requests.get(url)
        response.raise_for_status()
        content_type = response.headers['content-type'].split(';')[0]
        return response.content, content_type
    except requests.exceptions.RequestException as e:
        print(f"Error fetching document: {e}")
        return None, None

def on_chunk_select(event):
    selection = event.widget.curselection()
    if selection:
        index = selection[0]
        chunk_text_widget.delete(1.0, tk.END)
        prefix_text = prefix_text_widget.get(1.0, tk.END).strip() + "\n"
        chunk_text_widget.insert(tk.END, prefix_text + chunk_texts[index])
        
def copy_to_clipboard():
    selection = chunk_listbox.curselection()
    if selection:
        index = selection[0]
        root.clipboard_clear()
        prefix_text = prefix_text_widget.get(1.0, tk.END).strip() + "\n"
        root.clipboard_append(prefix_text + chunk_texts[index])

root = tk.Tk()
root.title("GPT-4 Input Assistant")

chunk_size_label = tk.Label(root, text="Amount of characters to parse:\nRecommended: 12000 - 16000")
chunk_size_label.grid(row=0, column=0, sticky="ns")
chunk_size_entry = tk.Entry(root)
chunk_size_entry.insert(0, "15000")
chunk_size_entry.grid(row=1, column=0, sticky="ns")

prefix_label = tk.Label(root, text="Prefix:")
prefix_label.grid(row=2, column=0, sticky="ns")
prefix_frame = tk.Frame(root)
prefix_frame.grid(row=3, column=0, sticky="ns")

url_label = tk.Label(root, text="URL:")
url_label.grid(row=0, column=1, sticky="nsew")
url_entry = tk.Entry(root)
url_entry.grid(row=1, column=1, sticky="nsew")

read_url_button = tk.Button(root, text="Read from URL", command=open_url)
read_url_button.grid(row=2, column=1, sticky="w")

prefix_text_widget = tk.Text(prefix_frame, wrap=tk.WORD, height=3, width=40)
prefix_text_widget.pack(side=tk.LEFT, expand=True, fill=tk.X)
prefix_scrollbar = ttk.Scrollbar(prefix_frame, orient="vertical", command=prefix_text_widget.yview)
prefix_scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
prefix_text_widget.configure(yscrollcommand=prefix_scrollbar.set)

open_file_button = tk.Button(root, text="Open File", command=open_file)
open_file_button.grid(row=3, column=1, sticky="w")
copy_button = tk.Button(root, text="Copy to Clipboard", command=copy_to_clipboard)
copy_button.grid(row=4, column=1, sticky="w")

chunks_frame = tk.Frame(root)
chunks_frame.grid(row=5, column=0, columnspan=2, sticky="nsew")

chunk_listbox = tk.Listbox(chunks_frame)
chunk_listbox.pack(side=tk.LEFT, expand=False, fill=tk.Y, padx=(0, 5))
chunk_listbox.bind("<<ListboxSelect>>", on_chunk_select)

chunk_texts = []

chunk_listbox_scrollbar = ttk.Scrollbar(chunks_frame, orient="vertical", command=chunk_listbox.yview)
chunk_listbox_scrollbar.pack(side=tk.LEFT, fill=tk.Y)
chunk_listbox.configure(yscrollcommand=chunk_listbox_scrollbar.set)

chunk_text_widget = tk.Text(chunks_frame, wrap=tk.WORD)
chunk_text_widget.pack(side=tk.RIGHT, expand=True, fill=tk.BOTH)

chunk_text_scrollbar = ttk.Scrollbar(chunks_frame, orient="vertical", command=chunk_text_widget.yview)
chunk_text_scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
chunk_text_widget.configure(yscrollcommand=chunk_text_scrollbar.set)

root.grid_rowconfigure(5, weight=1)
root.grid_columnconfigure(0, weight=1)
root.grid_columnconfigure(1, weight=1)

root.mainloop()
